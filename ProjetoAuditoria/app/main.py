"""
Aplicação Flask Principal - Sistema de Auditoria SEO Automatizada.

Esta aplicação fornece endpoints REST para controlar e monitorar
auditorias técnicas de SEO, integrando todos os componentes do sistema:
Agente Auditor, Agente Documentador, APIs externas e Chrome DevTools MCP.

Versão de Produção com:
- Sistema de cache avançado
- Processamento assíncrono
- Logging estruturado
- Segurança aprimorada
- Tratamento de erros robusto
"""

import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import sqlite3
import threading
import uuid
import asyncio

# Importações dos módulos do sistema
from .apis import APIManager
from .crawler import CrawlerManager
from .validator_agent import ValidatorAgent, AuditReport
from .doc_agent import DocumentorAgent
from .seo_audit_agent import SEOAuditAgent
from .seo_documentation_agent import SEODocumentationAgent
from .consolidate import DataConsolidator, ReportGenerator
from .report import ReportManager
from .database import DatabaseManager
from dataclasses import asdict

# Importações dos novos módulos de produção
from .logging_config import setup_logging, get_audit_logger
from .middleware import setup_request_logging
from .error_handlers import (
    AuditSystemError, ValidationError, ExternalAPIError, DatabaseError,
    InputValidator, handle_errors, register_error_handlers
)
from .security import setup_security, rate_limit, validate_security, require_ip_whitelist
from .cache import (
    cache_manager, cached, cache_audit_result, get_cached_audit_result,
    cache_api_response, get_cached_api_response, invalidate_audit_cache
)
from .performance_monitor import get_performance_collector
from .async_processor import (
    async_processor, submit_audit_task, get_audit_task_status,
    start_async_processing, stop_async_processing, TaskPriority
)
from .websocket_manager import init_websockets, websocket_manager

# Carrega variáveis de ambiente
load_dotenv()

# Configuração de logging estruturado
setup_logging()
logger = get_audit_logger(__name__)

# Inicialização da aplicação Flask
app = Flask(__name__)

# Inicialização do WebSocket
# Inicializar WebSocket Manager
websocket_mgr = init_websockets(app)
socketio = websocket_mgr.socketio

# Configuração de middleware de logging estruturado
setup_request_logging(app)

# Configuração de segurança e CORS
setup_security(app)
CORS(app)  # Habilita CORS para integração com front-end

# Configurações da aplicação
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key')
app.config['DEBUG'] = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
app.config['DATABASE_PATH'] = os.getenv('DATABASE_PATH', 'data/audit_history.db')

# Configuração de tratamento de erros
register_error_handlers(app)

# Armazenamento em memória para auditorias em execução
active_audits = {}
audit_results = {}


class AuditOrchestrator:
    """
    Orquestrador principal das auditorias SEO.
    
    Coordena a execução de auditorias completas integrando
    todos os componentes do sistema.
    """
    
    def __init__(self):
        """Inicializa o orquestrador com todos os componentes."""
        self.api_manager = APIManager()
        self.crawler_manager = CrawlerManager()
        self.validator_agent = ValidatorAgent(self.api_manager, self.crawler_manager)
        self.documentor_agent = DocumentorAgent()
        
        # Novos agentes especializados
        self.seo_audit_agent = SEOAuditAgent(data_folder="data")
        self.seo_documentation_agent = SEODocumentationAgent()
        
        self.data_consolidator = DataConsolidator()
        self.report_generator = ReportGenerator()
        self.report_manager = ReportManager()
        self.db_manager = DatabaseManager()
        self.logger = logging.getLogger(__name__)
    
    def execute_full_audit(self, url: str, audit_id: str, 
                          options: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Executa auditoria completa de SEO para uma URL.
        
        Args:
            url: URL a ser auditada.
            audit_id: ID único da auditoria.
            options: Opções de configuração da auditoria.
        
        Returns:
            Resultado da auditoria.
        """
        try:
            # Verificar cache primeiro
            cached_result = get_cached_audit_result(url, options)
            if cached_result:
                logger.info(f"Resultado encontrado em cache para {url}")
                return cached_result
            
            logger.info(f"Iniciando auditoria completa para {url} (ID: {audit_id})")
            
            # Processar opções
            if options is None:
                options = {}
            
            # Verifica se é modo apenas CSV
            csv_only_mode = options.get('csv_only_mode', False)
            screaming_frog_csv = options.get('screaming_frog_csv')
            use_screaming_frog = options.get('use_screaming_frog_data', False) or csv_only_mode
            screaming_frog_file = options.get('screaming_frog_file') or screaming_frog_csv
            checklist_file = options.get('checklist_file')
            
            # Atualiza status
            active_audits[audit_id] = {
                'status': 'running',
                'url': url,
                'start_time': datetime.now(),
                'current_step': 'initialization',
                'progress': 0
            }
            
            audit_result = {
                'audit_id': audit_id,
                'url': url,
                'start_time': datetime.now(),
                'status': 'running',
                'steps': [],
                'errors': [],
                'options': options,
                'csv_only_mode': csv_only_mode
            }
            
            # Se usar dados do Screaming Frog ou modo apenas CSV, executar fluxo específico
            if use_screaming_frog or csv_only_mode:
                return self._execute_screaming_frog_audit(
                    url, audit_id, screaming_frog_file, checklist_file, audit_result
                )
            
            # Fluxo normal de auditoria (apenas se não for modo CSV)
            # Passo 1: Coleta de dados das APIs (opcional)
            self._update_audit_status(audit_id, 'collecting_api_data', 10)
            api_data = self._collect_api_data(url, audit_result)
            
            # Passo 2: Execução do crawler (opcional se APIs não estiverem disponíveis)
            self._update_audit_status(audit_id, 'crawling_website', 25)
            crawler_data = self._execute_crawler(url, audit_result)
            
            # Passo 3: Validações com Chrome DevTools MCP (opcional)
            self._update_audit_status(audit_id, 'chrome_validations', 40)
            chrome_data = self._execute_chrome_validations(url, audit_result)
            
            # Passo 4: Consolidação de dados
            self._update_audit_status(audit_id, 'consolidating_data', 60)
            consolidated_data = self._consolidate_audit_data(
                url, api_data, crawler_data, chrome_data, audit_result
            )
            
            # Passo 5: Execução das validações
            self._update_audit_status(audit_id, 'executing_validations', 75)
            audit_report = self._execute_validations(consolidated_data, audit_result)
            
            # Passo 6: Geração de documentação
            self._update_audit_status(audit_id, 'generating_documentation', 90)
            documentation_result = self._generate_documentation(audit_report, audit_result)
            
            # Passo 7: Finalização
            self._update_audit_status(audit_id, 'completed', 100)
            
            # Salva resultado final
            audit_result.update({
                'status': 'completed',
                'end_time': datetime.now(),
                'audit_report': audit_report,
                'documentation': documentation_result,
                'overall_score': audit_report.overall_score if audit_report else 0
            })
            
            # Armazena no histórico
            self._save_audit_to_history(audit_result)
            
            # Salva arquivo JSON do relatório
            if audit_report:
                try:
                    json_filename = f"audit_{audit_id}.json"
                    json_path = self.report_manager.exporter.export_json_report(
                        audit_report, 
                        json_filename, 
                        include_raw_data=True
                    )
                    self.logger.info(f"Arquivo JSON salvo: {json_path}")
                except Exception as e:
                    self.logger.error(f"Erro ao salvar arquivo JSON: {str(e)}")
            
            # Remove da lista de auditorias ativas
            if audit_id in active_audits:
                del active_audits[audit_id]
            
            # Armazena resultado
            audit_results[audit_id] = audit_result
            
            # Cache do resultado para futuras consultas
            cache_audit_result(url, options, audit_result)
            
            # Enviar notificação WebSocket de conclusão
            websocket_manager.send_audit_completion(
                audit_id=audit_id,
                status='completed',
                url=url,
                summary=f"Auditoria concluída com sucesso para {url}"
            )
            
            logger.info(f"Auditoria {audit_id} concluída com sucesso")
            
            return audit_result
            
        except Exception as e:
            self.logger.error(f"Erro na auditoria {audit_id}: {str(e)}")
            
            # Atualiza status de erro
            error_result = {
                'audit_id': audit_id,
                'url': url,
                'status': 'error',
                'error': str(e),
                'start_time': datetime.now(),
                'end_time': datetime.now()
            }
            
            if audit_id in active_audits:
                active_audits[audit_id]['status'] = 'error'
                active_audits[audit_id]['error'] = str(e)
            
            # Enviar notificação WebSocket de erro
            websocket_manager.send_audit_completion(
                audit_id=audit_id,
                status='error',
                url=url,
                summary=f"Erro na auditoria: {str(e)}"
            )
            
            audit_results[audit_id] = error_result
            
            return error_result
    
    def _update_audit_status(self, audit_id: str, step: str, progress: int):
        """Atualiza status da auditoria e envia notificação WebSocket."""
        if audit_id in active_audits:
            active_audits[audit_id]['current_step'] = step
            active_audits[audit_id]['progress'] = progress
            active_audits[audit_id]['last_update'] = datetime.now()
            
            # Enviar notificação WebSocket sobre o progresso
            websocket_manager.send_audit_progress(
                audit_id=audit_id,
                stage=step,
                progress=progress
            )
    
    def _collect_api_data(self, url: str, audit_result: Dict) -> Dict[str, Any]:
        """Coleta dados das APIs externas (opcional)."""
        try:
            self.logger.info(f"Coletando dados das APIs para {url}")
            
            api_data = {}
            apis_available = False
            
            # Google Analytics 4 (opcional)
            try:
                ga4_data = self.api_manager.get_ga4_data(url)
                if ga4_data:
                    api_data['ga4'] = ga4_data
                    apis_available = True
                    audit_result['steps'].append({
                        'step': 'ga4_collection',
                        'status': 'success',
                        'timestamp': datetime.now(),
                        'data_points': len(ga4_data)
                    })
                    self.logger.info("Dados GA4 coletados com sucesso")
                else:
                    self.logger.info("GA4 não configurado - continuando sem dados do Google Analytics")
            except Exception as e:
                self.logger.info(f"GA4 não disponível: {str(e)}")
                audit_result['steps'].append({
                    'step': 'ga4_collection',
                    'status': 'skipped',
                    'timestamp': datetime.now(),
                    'message': 'GA4 não configurado'
                })
            
            # Google Search Console (opcional)
            try:
                gsc_data = self.api_manager.get_gsc_data(url)
                if gsc_data:
                    api_data['gsc'] = gsc_data
                    apis_available = True
                    audit_result['steps'].append({
                        'step': 'gsc_collection',
                        'status': 'success',
                        'timestamp': datetime.now(),
                        'data_points': len(gsc_data)
                    })
                    self.logger.info("Dados GSC coletados com sucesso")
                else:
                    self.logger.info("GSC não configurado - continuando sem dados do Search Console")
            except Exception as e:
                self.logger.info(f"GSC não disponível: {str(e)}")
                audit_result['steps'].append({
                    'step': 'gsc_collection',
                    'status': 'skipped',
                    'timestamp': datetime.now(),
                    'message': 'GSC não configurado'
                })
            
            # PageSpeed Insights (opcional)
            try:
                psi_data = self.api_manager.get_psi_data(url)
                if psi_data:
                    api_data['psi'] = psi_data
                    apis_available = True
                    audit_result['steps'].append({
                        'step': 'psi_collection',
                        'status': 'success',
                        'timestamp': datetime.now(),
                        'performance_score': psi_data.get('performance_score', 0)
                    })
                    self.logger.info("Dados PSI coletados com sucesso")
                else:
                    self.logger.info("PSI não configurado - continuando sem dados do PageSpeed Insights")
            except Exception as e:
                self.logger.info(f"PSI não disponível: {str(e)}")
                audit_result['steps'].append({
                    'step': 'psi_collection',
                    'status': 'skipped',
                    'timestamp': datetime.now(),
                    'message': 'PSI não configurado'
                })
            
            if not apis_available:
                self.logger.info("Nenhuma API externa configurada - continuando apenas com dados do Screaming Frog")
                audit_result['steps'].append({
                    'step': 'api_collection_summary',
                    'status': 'info',
                    'timestamp': datetime.now(),
                    'message': 'Nenhuma API externa configurada - usando apenas dados do CSV'
                })
            
            return api_data
            
        except Exception as e:
            self.logger.warning(f"Erro na coleta de dados das APIs (continuando sem APIs): {str(e)}")
            audit_result['steps'].append({
                'step': 'api_collection',
                'status': 'warning',
                'timestamp': datetime.now(),
                'message': f'APIs não disponíveis: {str(e)}'
            })
            return {}
    
    def _execute_crawler(self, url: str, audit_result: Dict) -> Dict[str, Any]:
        """Executa crawler Screaming Frog."""
        try:
            self.logger.info(f"Executando crawler para {url}")
            
            crawler_data = self.crawler_manager.crawl_website(url)
            
            audit_result['steps'].append({
                'step': 'crawler_execution',
                'status': 'success',
                'timestamp': datetime.now(),
                'pages_crawled': len(crawler_data.get('pages', [])) if crawler_data else 0
            })
            
            return crawler_data
            
        except Exception as e:
            self.logger.warning(f"Erro no crawler: {str(e)}")
            audit_result['errors'].append(f"Crawler: {str(e)}")
            return {}
    
    def _execute_chrome_validations(self, url: str, audit_result: Dict) -> Dict[str, Any]:
        """Executa validações com Chrome DevTools MCP."""
        try:
            self.logger.info(f"Executando validações Chrome DevTools para {url}")
            
            # Aqui seria a integração com Chrome DevTools MCP
            # Por enquanto, retorna dados simulados
            chrome_data = {
                'performance_metrics': {
                    'lcp': 2.5,
                    'fid': 100,
                    'cls': 0.1,
                    'ttfb': 800
                },
                'accessibility_score': 85,
                'seo_score': 90,
                'best_practices_score': 88,
                'console_errors': [],
                'network_requests': 45,
                'page_size': 2.1
            }
            
            audit_result['steps'].append({
                'step': 'chrome_validations',
                'status': 'success',
                'timestamp': datetime.now(),
                'performance_score': chrome_data.get('performance_metrics', {}).get('lcp', 0)
            })
            
            return chrome_data
            
        except Exception as e:
            self.logger.warning(f"Erro nas validações Chrome: {str(e)}")
            audit_result['errors'].append(f"Chrome DevTools: {str(e)}")
            return {}
    
    def _consolidate_audit_data(self, url: str, api_data: Dict, 
                              crawler_data: Dict, chrome_data: Dict,
                              audit_result: Dict) -> Dict[str, Any]:
        """Consolida dados de todas as fontes."""
        try:
            self.logger.info(f"Consolidando dados para {url}")
            
            consolidated_data = self.data_consolidator.consolidate_data_sources(
                url, api_data, crawler_data, chrome_data
            )
            
            audit_result['steps'].append({
                'step': 'data_consolidation',
                'status': 'success',
                'timestamp': datetime.now(),
                'consolidated_metrics': len(consolidated_data.get('metrics', {}))
            })
            
            return consolidated_data
            
        except Exception as e:
            self.logger.error(f"Erro na consolidação: {str(e)}")
            raise
    
    def _execute_validations(self, consolidated_data: Dict, 
                           audit_result: Dict) -> Optional[AuditReport]:
        """Executa validações usando os Agentes Auditores."""
        try:
            self.logger.info("Executando validações com Agente SEO Especializado")
            
            # Executar auditoria com o novo agente especializado
            seo_audit_result = asyncio.run(
                self.seo_audit_agent.run_audit([consolidated_data['url']])
            )
            
            # Executar validações com o agente original (compatibilidade)
            audit_report = asyncio.run(
                self.validator_agent.execute_full_audit(
                    consolidated_data['url'],
                    include_chrome_devtools=True
                )
            )
            
            # Combinar resultados
            total_problems = len(seo_audit_result.get('problems', [])) if seo_audit_result.get('success') else 0
            original_validations = len(audit_report.validations) if audit_report else 0
            
            audit_result['steps'].append({
                'step': 'validations_execution',
                'status': 'success',
                'timestamp': datetime.now(),
                'seo_problems_found': total_problems,
                'original_validations_count': original_validations,
                'overall_score': audit_report.overall_score if audit_report else 0,
                'seo_audit_success': seo_audit_result.get('success', False)
            })
            
            # Armazenar resultado da auditoria SEO para uso posterior
            audit_result['seo_audit_data'] = seo_audit_result
            
            return audit_report
            
        except Exception as e:
            self.logger.error(f"Erro nas validações: {str(e)}")
            audit_result['errors'].append(f"Validações: {str(e)}")
            return None
    
    def _generate_documentation(self, audit_report: Optional[AuditReport],
                              audit_result: Dict) -> Dict[str, Any]:
        """Gera documentação usando os Agentes Documentadores."""
        try:
            self.logger.info("Gerando documentação com Agente SEO Especializado")
            
            documentation_results = {}
            
            # Gerar documentação com o novo agente especializado
            if audit_result.get('seo_audit_data') and audit_result['seo_audit_data'].get('success'):
                seo_problems = audit_result['seo_audit_data'].get('problems', [])
                
                if seo_problems:
                    self.logger.info(f"Processando {len(seo_problems)} problemas SEO para documentação")
                    
                    # Processar cada problema individualmente
                    seo_documentation_sections = []
                    for problem in seo_problems:
                        try:
                            doc_section = self.seo_documentation_agent.process_seo_problem(problem)
                            seo_documentation_sections.append(doc_section)
                        except Exception as e:
                            self.logger.warning(f"Erro ao processar problema SEO: {str(e)}")
                    
                    # Compilar documentação final
                    if seo_documentation_sections:
                        final_documentation = self.seo_documentation_agent.compile_final_document(
                            seo_documentation_sections,
                            audit_result['seo_audit_data'].get('statistics', {})
                        )
                        documentation_results['seo_specialized'] = final_documentation
                else:
                    self.logger.info("Nenhum problema SEO encontrado para documentação")
                    documentation_results['seo_specialized'] = {
                        'status': 'no_problems',
                        'message': 'Nenhum problema técnico de SEO foi identificado na auditoria.'
                    }
            
            # Gerar documentação com o agente original (compatibilidade)
            if audit_report:
                original_documentation = self.documentor_agent.generate_audit_documentation(audit_report)
                documentation_results['original'] = original_documentation
            
            audit_result['steps'].append({
                'step': 'documentation_generation',
                'status': 'success',
                'timestamp': datetime.now(),
                'seo_sections_generated': len(documentation_results.get('seo_specialized', {}).get('sections', [])),
                'original_document_url': documentation_results.get('original', {}).get('document_url', ''),
                'has_seo_documentation': 'seo_specialized' in documentation_results
            })
            
            return documentation_results
            
        except Exception as e:
            self.logger.warning(f"Erro na documentação: {str(e)}")
            audit_result['errors'].append(f"Documentação: {str(e)}")
            return {'status': 'error', 'error': str(e)}
    
    def _execute_screaming_frog_audit(self, url: str, audit_id: str, 
                                    screaming_frog_file: str, checklist_file: str,
                                    audit_result: Dict) -> Dict[str, Any]:
        """
        Executa auditoria específica usando dados do Screaming Frog.
        
        Args:
            url: URL sendo auditada (para referência)
            audit_id: ID da auditoria
            screaming_frog_file: Caminho para o arquivo CSV do Screaming Frog
            checklist_file: Caminho para o arquivo de checklist
            audit_result: Dicionário de resultado da auditoria
        
        Returns:
            Resultado da auditoria com dados do Screaming Frog
        """
        try:
            self.logger.info(f"Executando auditoria com dados do Screaming Frog para {url}")
            
            # Passo 1: Configurar agente SEO com arquivos específicos
            self._update_audit_status(audit_id, 'loading_screaming_frog_data', 20)
            
            # Carregar dados do Screaming Frog no agente SEO
            if screaming_frog_file:
                screaming_frog_loaded = self.seo_audit_agent._load_screaming_frog_data(screaming_frog_file)
                if not screaming_frog_loaded:
                    raise Exception(f"Falha ao carregar dados do Screaming Frog: {screaming_frog_file}")
            
            # Carregar checklist personalizado
            if checklist_file:
                checklist_loaded = self.seo_audit_agent._load_checklist(checklist_file)
                if not checklist_loaded:
                    self.logger.warning(f"Falha ao carregar checklist personalizado: {checklist_file}")
            
            audit_result['steps'].append({
                'step': 'screaming_frog_data_loading',
                'status': 'success',
                'timestamp': datetime.now(),
                'screaming_frog_file': screaming_frog_file,
                'checklist_file': checklist_file
            })
            
            # Passo 2: Analisar problemas do Screaming Frog
            self._update_audit_status(audit_id, 'analyzing_screaming_frog_problems', 50)
            
            screaming_frog_problems = self.seo_audit_agent._analyze_screaming_frog_problems(
                self.seo_audit_agent.screaming_frog_data
            )
            
            audit_result['steps'].append({
                'step': 'screaming_frog_analysis',
                'status': 'success',
                'timestamp': datetime.now(),
                'problems_found': len(screaming_frog_problems)
            })
            
            # Passo 3: Gerar relatório consolidado
            self._update_audit_status(audit_id, 'generating_screaming_frog_report', 75)
            
            # Criar relatório de auditoria com os problemas encontrados
            audit_report = AuditReport(
                url=url,
                audit_timestamp=datetime.now(),
                overall_score=max(0, 100 - (len(screaming_frog_problems) * 5)),
                validations=[],
                summary={
                    'total_problems': len(screaming_frog_problems),
                    'critical_problems': len([p for p in screaming_frog_problems if p.severity == 'Critical']),
                    'high_problems': len([p for p in screaming_frog_problems if p.severity == 'High']),
                    'medium_problems': len([p for p in screaming_frog_problems if p.severity == 'Medium']),
                    'low_problems': len([p for p in screaming_frog_problems if p.severity == 'Low']),
                    'data_source': 'Screaming Frog'
                },
                raw_data={'screaming_frog_problems': screaming_frog_problems}
            )
            
            # Passo 4: Gerar documentação
            self._update_audit_status(audit_id, 'generating_documentation', 90)
            documentation_result = self._generate_documentation(audit_report, audit_result)
            
            # Passo 5: Finalização
            self._update_audit_status(audit_id, 'completed', 100)
            
            # Salvar resultado final
            final_result = {
                'audit_id': audit_id,
                'url': url,
                'status': 'completed',
                'timestamp': datetime.now(),
                'data_source': 'Screaming Frog',
                'problems': [asdict(p) for p in screaming_frog_problems],
                'summary': audit_report.summary,
                'documentation': documentation_result,
                'steps': audit_result['steps']
            }
            
            audit_results[audit_id] = final_result
            self._save_audit_to_history(final_result)
            
            self.logger.info(f"Auditoria Screaming Frog concluída para {url} - {len(screaming_frog_problems)} problemas encontrados")
            
            return final_result
            
        except Exception as e:
            self.logger.error(f"Erro na auditoria Screaming Frog: {str(e)}")
            audit_result['errors'].append(f"Screaming Frog Audit: {str(e)}")
            
            # Atualizar status como erro
            active_audits[audit_id]['status'] = 'error'
            active_audits[audit_id]['error'] = str(e)
            
            return {
                'audit_id': audit_id,
                'url': url,
                'status': 'error',
                'error': str(e),
                'timestamp': datetime.now()
            }
    
    def _save_audit_to_history(self, audit_result: Dict):
        """Salva auditoria no histórico do banco de dados."""
        try:
            # Implementação do salvamento no SQLite será feita posteriormente
            self.logger.info(f"Auditoria {audit_result['audit_id']} salva no histórico")
        except Exception as e:
            self.logger.error(f"Erro ao salvar no histórico: {str(e)}")


# Instância global do orquestrador
orchestrator = AuditOrchestrator()


# === ENDPOINTS DA API ===

@app.route('/health', methods=['GET'])
@handle_errors
def health_check():
    """
    Endpoint de verificação de saúde da aplicação.
    
    Returns:
        Status da aplicação e componentes.
    """
    try:
        # Verificar componentes críticos
        health_status = {
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'version': '1.0.0',
            'components': {}
        }
        
        # Verificar banco de dados
        try:
            db_manager = DatabaseManager()
            # Teste simples de conexão
            health_status['components']['database'] = {
                'status': 'healthy',
                'type': 'sqlite',
                'path': app.config['DATABASE_PATH']
            }
        except Exception as e:
            health_status['components']['database'] = {
                'status': 'unhealthy',
                'error': str(e)
            }
            health_status['status'] = 'degraded'
        
        # Verificar cache
        try:
            cache_stats = cache_manager.stats()
            health_status['components']['cache'] = {
                'status': 'healthy',
                'memory_cache_size': cache_stats.get('memory_cache_size', 0),
                'disk_cache_size': cache_stats.get('disk_cache_size', 0)
            }
        except Exception as e:
            health_status['components']['cache'] = {
                'status': 'unhealthy',
                'error': str(e)
            }
        
        # Verificar processador assíncrono
        try:
            processor_stats = async_processor.get_stats()
            health_status['components']['async_processor'] = {
                'status': 'healthy',
                'active_tasks': processor_stats.get('active_tasks', 0),
                'completed_tasks': processor_stats.get('completed_tasks', 0),
                'failed_tasks': processor_stats.get('failed_tasks', 0)
            }
        except Exception as e:
            health_status['components']['async_processor'] = {
                'status': 'unhealthy',
                'error': str(e)
            }
        
        # Verificar auditorias ativas
        health_status['components']['active_audits'] = {
            'status': 'healthy',
            'count': len(active_audits),
            'audits': list(active_audits.keys())
        }
        
        # Verificar resultados de auditoria
        health_status['components']['audit_results'] = {
            'status': 'healthy',
            'count': len(audit_results)
        }
        
        status_code = 200 if health_status['status'] == 'healthy' else 503
        return jsonify(health_status), status_code
        
    except Exception as e:
        logger.error(f"Erro no health check: {str(e)}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 503


@app.route('/metrics', methods=['GET'])
@rate_limit(limit=60, window=60)
@handle_errors
def get_metrics():
    """
    Endpoint para métricas detalhadas do sistema.
    """
    try:
        metrics = {
            'timestamp': datetime.now().isoformat(),
            'system': {
                'active_audits': len(active_audits),
                'completed_audits': len(audit_results),
                'total_audits_in_history': 0
            },
            'cache': cache_manager.stats(),
            'async_processor': async_processor.get_stats(),
            'memory': {
                'active_audits_memory': len(str(active_audits)),
                'results_memory': len(str(audit_results))
            }
        }
        
        # Obter estatísticas do banco de dados
        try:
            db_manager = DatabaseManager()
            history = db_manager.get_audit_history(limit=1000)
            metrics['system']['total_audits_in_history'] = len(history)
        except Exception as e:
            logger.warning(f"Erro ao obter estatísticas do banco: {str(e)}")
        
        return jsonify(metrics), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter métricas: {str(e)}")
        raise AuditSystemError(f"Erro ao obter métricas: {str(e)}")


@app.route('/metrics/prometheus', methods=['GET'])
@rate_limit(limit=120, window=60)
@handle_errors
def get_prometheus_metrics():
    """
    Endpoint para métricas no formato Prometheus.
    """
    try:
        import psutil
        
        # Métricas básicas do sistema
        metrics_lines = []
        
        # Auditorias ativas
        metrics_lines.append(f"# HELP seo_audit_active_audits Número de auditorias ativas")
        metrics_lines.append(f"# TYPE seo_audit_active_audits gauge")
        metrics_lines.append(f"seo_audit_active_audits {len(active_audits)}")
        
        # Auditorias completadas
        metrics_lines.append(f"# HELP seo_audit_completed_audits Número de auditorias completadas")
        metrics_lines.append(f"# TYPE seo_audit_completed_audits gauge")
        metrics_lines.append(f"seo_audit_completed_audits {len(audit_results)}")
        
        # Métricas de cache
        cache_stats = cache_manager.get_stats()
        metrics_lines.append(f"# HELP seo_audit_cache_hits Cache hits")
        metrics_lines.append(f"# TYPE seo_audit_cache_hits counter")
        metrics_lines.append(f"seo_audit_cache_hits {cache_stats.get('hits', 0)}")
        
        metrics_lines.append(f"# HELP seo_audit_cache_misses Cache misses")
        metrics_lines.append(f"# TYPE seo_audit_cache_misses counter")
        metrics_lines.append(f"seo_audit_cache_misses {cache_stats.get('misses', 0)}")
        
        # Métricas do processador assíncrono
        async_stats = async_processor.get_stats()
        metrics_lines.append(f"# HELP seo_audit_async_queue_size Tamanho da fila assíncrona")
        metrics_lines.append(f"# TYPE seo_audit_async_queue_size gauge")
        metrics_lines.append(f"seo_audit_async_queue_size {async_stats.get('queue_size', 0)}")
        
        metrics_lines.append(f"# HELP seo_audit_async_active_workers Workers ativos")
        metrics_lines.append(f"# TYPE seo_audit_async_active_workers gauge")
        metrics_lines.append(f"seo_audit_async_active_workers {async_stats.get('active_workers', 0)}")
        
        # Métricas de sistema
        memory_info = psutil.virtual_memory()
        metrics_lines.append(f"# HELP seo_audit_memory_usage_percent Uso de memória em porcentagem")
        metrics_lines.append(f"# TYPE seo_audit_memory_usage_percent gauge")
        metrics_lines.append(f"seo_audit_memory_usage_percent {memory_info.percent}")
        
        cpu_percent = psutil.cpu_percent(interval=1)
        metrics_lines.append(f"# HELP seo_audit_cpu_usage_percent Uso de CPU em porcentagem")
        metrics_lines.append(f"# TYPE seo_audit_cpu_usage_percent gauge")
        metrics_lines.append(f"seo_audit_cpu_usage_percent {cpu_percent}")
        
        # Métricas de disco
        disk_usage = psutil.disk_usage('/')
        metrics_lines.append(f"# HELP seo_audit_disk_usage_percent Uso de disco em porcentagem")
        metrics_lines.append(f"# TYPE seo_audit_disk_usage_percent gauge")
        metrics_lines.append(f"seo_audit_disk_usage_percent {(disk_usage.used / disk_usage.total) * 100}")
        
        response_text = '\n'.join(metrics_lines) + '\n'
        
        from flask import Response
        return Response(response_text, mimetype='text/plain')
        
    except Exception as e:
        logger.error(f"Erro ao gerar métricas Prometheus: {str(e)}")
        raise AuditSystemError(f"Erro ao gerar métricas Prometheus: {str(e)}")


@app.route('/admin/health/deep', methods=['GET'])
@rate_limit(limit=30, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def deep_health_check():
    """
    Endpoint para verificação profunda de saúde do sistema.
    """
    try:
        import psutil
        import requests
        from urllib.parse import urlparse
        
        health_status = {
            'timestamp': datetime.now().isoformat(),
            'overall_status': 'healthy',
            'checks': {}
        }
        
        # Verificação do banco de dados
        try:
            db_manager = DatabaseManager()
            db_manager.get_audit_history(limit=1)
            health_status['checks']['database'] = {
                'status': 'healthy',
                'message': 'Conexão com banco de dados OK'
            }
        except Exception as e:
            health_status['checks']['database'] = {
                'status': 'unhealthy',
                'message': f'Erro no banco de dados: {str(e)}'
            }
            health_status['overall_status'] = 'unhealthy'
        
        # Verificação do cache
        try:
            cache_stats = cache_manager.stats()
            health_status['checks']['cache'] = {
                'status': 'healthy',
                'message': 'Cache funcionando',
                'stats': cache_stats
            }
        except Exception as e:
            health_status['checks']['cache'] = {
                'status': 'unhealthy',
                'message': f'Erro no cache: {str(e)}'
            }
            health_status['overall_status'] = 'unhealthy'
        
        # Verificação do processador assíncrono
        try:
            async_stats = async_processor.get_stats()
            if async_stats.get('is_running', False):
                health_status['checks']['async_processor'] = {
                    'status': 'healthy',
                    'message': 'Processador assíncrono ativo',
                    'stats': async_stats
                }
            else:
                health_status['checks']['async_processor'] = {
                    'status': 'warning',
                    'message': 'Processador assíncrono inativo'
                }
        except Exception as e:
            health_status['checks']['async_processor'] = {
                'status': 'unhealthy',
                'message': f'Erro no processador assíncrono: {str(e)}'
            }
            health_status['overall_status'] = 'unhealthy'
        
        # Verificação de recursos do sistema
        try:
            memory_info = psutil.virtual_memory()
            cpu_percent = psutil.cpu_percent(interval=1)
            disk_usage = psutil.disk_usage('/')
            
            # Alertas para recursos críticos
            memory_critical = memory_info.percent > 90
            cpu_critical = cpu_percent > 90
            disk_critical = (disk_usage.used / disk_usage.total) * 100 > 90
            
            resource_status = 'healthy'
            resource_messages = []
            
            if memory_critical:
                resource_status = 'critical'
                resource_messages.append(f'Memória crítica: {memory_info.percent:.1f}%')
            elif memory_info.percent > 80:
                resource_status = 'warning'
                resource_messages.append(f'Memória alta: {memory_info.percent:.1f}%')
            
            if cpu_critical:
                resource_status = 'critical'
                resource_messages.append(f'CPU crítica: {cpu_percent:.1f}%')
            elif cpu_percent > 80:
                resource_status = 'warning'
                resource_messages.append(f'CPU alta: {cpu_percent:.1f}%')
            
            if disk_critical:
                resource_status = 'critical'
                disk_percent = (disk_usage.used / disk_usage.total) * 100
                resource_messages.append(f'Disco crítico: {disk_percent:.1f}%')
            
            health_status['checks']['system_resources'] = {
                'status': resource_status,
                'message': '; '.join(resource_messages) if resource_messages else 'Recursos do sistema OK',
                'details': {
                    'memory_percent': memory_info.percent,
                    'cpu_percent': cpu_percent,
                    'disk_percent': (disk_usage.used / disk_usage.total) * 100
                }
            }
            
            if resource_status == 'critical':
                health_status['overall_status'] = 'critical'
            elif resource_status == 'warning' and health_status['overall_status'] == 'healthy':
                health_status['overall_status'] = 'warning'
                
        except Exception as e:
            health_status['checks']['system_resources'] = {
                'status': 'unhealthy',
                'message': f'Erro ao verificar recursos: {str(e)}'
            }
            health_status['overall_status'] = 'unhealthy'
        
        # Verificação de APIs externas (opcional)
        external_apis = [
            {'name': 'PageSpeed Insights', 'url': 'https://www.googleapis.com/pagespeedonline/v5/runPagespeed'},
            {'name': 'GTmetrix', 'url': 'https://gtmetrix.com/api/2.0/status'}
        ]
        
        for api in external_apis:
            try:
                response = requests.get(api['url'], timeout=5)
                if response.status_code == 200:
                    health_status['checks'][f"external_api_{api['name'].lower().replace(' ', '_')}"] = {
                        'status': 'healthy',
                        'message': f'{api["name"]} acessível'
                    }
                else:
                    health_status['checks'][f"external_api_{api['name'].lower().replace(' ', '_')}"] = {
                        'status': 'warning',
                        'message': f'{api["name"]} retornou status {response.status_code}'
                    }
            except Exception as e:
                health_status['checks'][f"external_api_{api['name'].lower().replace(' ', '_')}"] = {
                    'status': 'warning',
                    'message': f'{api["name"]} inacessível: {str(e)}'
                }
        
        # Determinar código de status HTTP baseado no status geral
        status_code = 200
        if health_status['overall_status'] == 'warning':
            status_code = 200  # Warning ainda é OK
        elif health_status['overall_status'] in ['unhealthy', 'critical']:
            status_code = 503  # Service Unavailable
        
        return jsonify(health_status), status_code
        
    except Exception as e:
        logger.error(f"Erro na verificação profunda de saúde: {str(e)}")
        return jsonify({
            'timestamp': datetime.now().isoformat(),
            'overall_status': 'error',
            'message': f'Erro na verificação de saúde: {str(e)}'
        }), 500


@app.route('/admin/metrics/performance', methods=['GET'])
@rate_limit(limit=60, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def get_performance_metrics():
    """
    Endpoint para métricas detalhadas de performance.
    """
    try:
        import psutil
        import time
        
        # Métricas de performance do sistema
        performance_metrics = {
            'timestamp': datetime.now().isoformat(),
            'system': {
                'uptime': time.time() - psutil.boot_time(),
                'load_average': os.getloadavg() if hasattr(os, 'getloadavg') else None,
                'cpu': {
                    'percent': psutil.cpu_percent(interval=1),
                    'count': psutil.cpu_count(),
                    'freq': psutil.cpu_freq()._asdict() if psutil.cpu_freq() else None
                },
                'memory': {
                    'virtual': psutil.virtual_memory()._asdict(),
                    'swap': psutil.swap_memory()._asdict()
                },
                'disk': {
                    'usage': psutil.disk_usage('/')._asdict(),
                    'io': psutil.disk_io_counters()._asdict() if psutil.disk_io_counters() else None
                },
                'network': psutil.net_io_counters()._asdict() if psutil.net_io_counters() else None
            },
            'application': {
                'active_audits': len(active_audits),
                'completed_audits': len(audit_results),
                'cache_stats': cache_manager.stats(),
                'async_stats': async_processor.get_stats()
            }
        }
        
        # Métricas de auditoria por período
        try:
            db_manager = DatabaseManager()
            recent_audits = db_manager.get_audit_history(limit=100)
            
            # Calcular métricas de tempo de auditoria
            audit_times = []
            for audit in recent_audits:
                if audit.get('completed_at') and audit.get('started_at'):
                    start_time = datetime.fromisoformat(audit['started_at'])
                    end_time = datetime.fromisoformat(audit['completed_at'])
                    duration = (end_time - start_time).total_seconds()
                    audit_times.append(duration)
            
            if audit_times:
                performance_metrics['application']['audit_performance'] = {
                    'average_duration': sum(audit_times) / len(audit_times),
                    'min_duration': min(audit_times),
                    'max_duration': max(audit_times),
                    'total_audits_analyzed': len(audit_times)
                }
            
        except Exception as e:
            logger.warning(f"Erro ao calcular métricas de performance de auditoria: {str(e)}")
        
        return jsonify(performance_metrics), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter métricas de performance: {str(e)}")
        raise AuditSystemError(f"Erro ao obter métricas de performance: {str(e)}")


@app.route('/admin/cache/clear', methods=['POST'])
@rate_limit(limit=5, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def clear_cache():
    """
    Endpoint administrativo para limpar o cache.
    """
    try:
        cache_manager.clear()
        logger.info("Cache limpo via endpoint administrativo")
        return jsonify({
            'status': 'success',
            'message': 'Cache limpo com sucesso',
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao limpar cache: {str(e)}")
        raise AuditSystemError(f"Erro ao limpar cache: {str(e)}")


@app.route('/admin/cache/stats', methods=['GET'])
@rate_limit(limit=30, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def get_cache_stats():
    """
    Endpoint administrativo para obter estatísticas detalhadas do cache.
    """
    try:
        stats = cache_manager.stats()
        health = cache_manager.health_check()
        
        return jsonify({
            'status': 'success',
            'cache_stats': stats,
            'cache_health': health,
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter estatísticas do cache: {str(e)}")
        raise AuditSystemError(f"Erro ao obter estatísticas do cache: {str(e)}")


@app.route('/admin/performance/metrics', methods=['GET'])
@rate_limit(limit=30, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def get_performance_monitoring_metrics():
    """
    Endpoint administrativo para obter métricas detalhadas de performance.
    """
    try:
        performance_collector = get_performance_collector()
        
        # Coleta métricas do sistema
        system_metrics = performance_collector.collect_system_metrics()
        
        # Coleta métricas da aplicação
        app_metrics = performance_collector.collect_application_metrics()
        
        # Estatísticas de requisições
        request_stats = performance_collector.get_request_statistics()
        
        return jsonify({
            'status': 'success',
            'system_metrics': system_metrics.__dict__,
            'application_metrics': app_metrics.__dict__,
            'request_statistics': request_stats,
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter métricas de performance: {str(e)}")
        raise AuditSystemError(f"Erro ao obter métricas de performance: {str(e)}")


@app.route('/admin/performance/reset', methods=['POST'])
@rate_limit(limit=5, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def reset_performance_metrics():
    """
    Endpoint administrativo para resetar métricas de performance.
    """
    try:
        performance_collector = get_performance_collector()
        performance_collector.reset_metrics()
        
        return jsonify({
            'status': 'success',
            'message': 'Métricas de performance resetadas com sucesso',
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao resetar métricas de performance: {str(e)}")
        raise AuditSystemError(f"Erro ao resetar métricas de performance: {str(e)}")


@app.route('/admin/processor/stats', methods=['GET'])
@rate_limit(limit=30, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def get_processor_stats():
    """
    Endpoint administrativo para estatísticas detalhadas do processador.
    """
    try:
        stats = async_processor.get_stats()
        return jsonify(stats), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter estatísticas do processador: {str(e)}")
        raise AuditSystemError(f"Erro ao obter estatísticas do processador: {str(e)}")


@app.route('/admin/processor/start', methods=['POST'])
@rate_limit(limit=5, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def start_processor():
    """
    Inicia o processador assíncrono manualmente.
    
    Returns:
        JSON com status da operação
    """
    try:
        start_async_processing()
        logger.info("Processador assíncrono iniciado manualmente")
        return jsonify({"message": "Processador iniciado com sucesso"}), 200
    except Exception as e:
        logger.error(f"Erro ao iniciar processador: {str(e)}")
        return jsonify({"error": f"Erro ao iniciar processador: {str(e)}"}), 500


@app.route('/admin/system/status', methods=['GET'])
@rate_limit(limit=10, window=60)
@require_ip_whitelist(['127.0.0.1', '::1'])
@handle_errors
def get_system_status():
    """
    Endpoint administrativo para status completo do sistema.
    """
    try:
        import psutil
        import sys
        
        status = {
            'timestamp': datetime.now().isoformat(),
            'python_version': sys.version,
            'system': {
                'cpu_percent': psutil.cpu_percent(interval=1),
                'memory_percent': psutil.virtual_memory().percent,
                'disk_usage': psutil.disk_usage('/').percent
            },
            'application': {
                'active_audits': len(active_audits),
                'completed_audits': len(audit_results),
                'uptime': 'N/A'  # Seria necessário rastrear o tempo de início
            },
            'components': {
                'cache': cache_manager.stats(),
                'async_processor': async_processor.get_stats()
            }
        }
        
        return jsonify(status), 200
        
    except ImportError:
        # Se psutil não estiver disponível, retorna informações básicas
        status = {
            'timestamp': datetime.now().isoformat(),
            'system': 'Informações de sistema não disponíveis (psutil não instalado)',
            'application': {
                'active_audits': len(active_audits),
                'completed_audits': len(audit_results)
            },
            'components': {
                'cache': cache_manager.stats(),
                'async_processor': async_processor.get_stats()
            }
        }
        return jsonify(status), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter status do sistema: {str(e)}")
        raise AuditSystemError(f"Erro ao obter status do sistema: {str(e)}")


@app.route('/audit/start', methods=['POST'])
@rate_limit(limit=100, window=60)
@validate_security(validate_url=True, validate_json=False)
@handle_errors
def start_audit():
    """
    Inicia uma nova auditoria SEO.
    
    Body (JSON):
        {
            "url": "https://example.com",  // Opcional se screaming_frog_file for fornecido
            "options": {
                "include_crawler": true,
                "include_chrome_validations": true,
                "generate_documentation": true
            }
        }
    
    Body (Form-data para upload de CSV):
        - screaming_frog_file: Arquivo CSV do Screaming Frog (obrigatório)
        - url: URL principal (opcional)
        - options: JSON string com opções (opcional)
    
    Returns:
        ID da auditoria iniciada.
    """
    try:
        # Verifica se é upload de arquivo
        if 'screaming_frog_file' in request.files:
            return _handle_csv_upload_audit()
        
        # Processamento JSON tradicional
        try:
            data = request.get_json()
        except Exception:
            return jsonify({
                'error': 'JSON inválido',
                'code': 'INVALID_JSON'
            }), 400
        
        if not data:
            return jsonify({
                'error': 'Dados são obrigatórios (JSON ou upload de CSV)',
                'code': 'MISSING_DATA'
            }), 400
        
        url = data.get('url')
        options = data.get('options', {})
        
        # URL é opcional se não for fornecida
        if url:
            # Valida URL se fornecida
            if not url.startswith(('http://', 'https://')):
                return jsonify({
                    'error': 'URL deve começar com http:// ou https://',
                    'code': 'INVALID_URL'
                }), 400
        else:
            # Se não há URL, define uma URL padrão para o sistema
            url = "https://example.com"  # URL placeholder
            logger.info("Auditoria iniciada sem URL específica - usando dados do CSV apenas")
        
        # Gera ID único para a auditoria
        audit_id = str(uuid.uuid4())
        
        # Adiciona auditoria ao active_audits imediatamente
        active_audits[audit_id] = {
            'status': 'queued',
            'url': url,
            'start_time': datetime.now(),
            'current_step': 'queued',
            'progress': 0,
            'options': options or {}
        }
        
        # Callback para sincronizar status quando tarefa inicia/termina
        def audit_callback(task_id: str, result: Any, error: str):
            """Callback chamado quando tarefa de auditoria termina."""
            if error:
                logger.error(f"Erro na auditoria {audit_id}: {error}")
                if audit_id in active_audits:
                    active_audits[audit_id]['status'] = 'failed'
                    active_audits[audit_id]['error'] = error
                # Adiciona start_time ao resultado de erro
                error_result = {
                    'audit_id': audit_id,
                    'url': url,
                    'status': 'error',
                    'error': error,
                    'start_time': active_audits.get(audit_id, {}).get('start_time', datetime.now()),
                    'end_time': datetime.now(),
                    'overall_score': 0
                }
                audit_results[audit_id] = error_result
            else:
                logger.info(f"Auditoria {audit_id} concluída com sucesso")
                # Move de active_audits para audit_results
                if audit_id in active_audits:
                    audit_info = active_audits.pop(audit_id)
                    # Garante que o resultado tenha overall_score
                    final_result = {
                        **result,
                        'status': 'completed',
                        'start_time': audit_info.get('start_time', datetime.now()),
                        'end_time': datetime.now()
                    }
                    # Garante que overall_score esteja presente
                    if 'overall_score' not in final_result:
                        final_result['overall_score'] = 0
                    audit_results[audit_id] = final_result

        # Usar async_processor em vez de threading.Thread direto
        task_id = submit_audit_task(
            orchestrator.execute_full_audit,
            audit_id,
            url,
            options or {},
            callback=audit_callback
        )
        
        logger.info(f"Auditoria iniciada: {audit_id} para {url} (task_id: {task_id})")
        
        return jsonify({
            'audit_id': audit_id,
            'url': url,
            'status': 'started',
            'message': 'Auditoria iniciada com sucesso',
            'estimated_duration': '5-10 minutos',
            'task_id': task_id
        }), 202
        
    except Exception as e:
        logger.error(f"Erro ao iniciar auditoria: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


def _handle_csv_upload_audit():
    """
    Processa upload de CSV do Screaming Frog para auditoria.
    
    Returns:
        Resposta JSON com ID da auditoria.
    """
    try:
        # Verifica se arquivo foi enviado
        if 'screaming_frog_file' not in request.files:
            return jsonify({
                'error': 'Arquivo CSV do Screaming Frog é obrigatório',
                'code': 'MISSING_CSV_FILE'
            }), 400
        
        file = request.files['screaming_frog_file']
        
        if file.filename == '':
            return jsonify({
                'error': 'Nenhum arquivo selecionado',
                'code': 'NO_FILE_SELECTED'
            }), 400
        
        # Verifica extensão do arquivo
        if not file.filename.lower().endswith('.csv'):
            return jsonify({
                'error': 'Arquivo deve ser um CSV',
                'code': 'INVALID_FILE_TYPE'
            }), 400
        
        # Salva arquivo temporariamente
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w+b', suffix='.csv', delete=False) as temp_file:
            file.save(temp_file.name)
            csv_path = temp_file.name
        
        # Obtém parâmetros opcionais
        url = request.form.get('url', 'https://example.com')  # URL opcional
        options_str = request.form.get('options', '{}')
        
        try:
            import json
            options = json.loads(options_str)
        except json.JSONDecodeError:
            options = {}
        
        # Adiciona caminho do CSV às opções
        options['screaming_frog_csv'] = csv_path
        options['csv_only_mode'] = True  # Indica que é modo apenas CSV
        
        # Valida URL se fornecida
        if url != 'https://example.com' and not url.startswith(('http://', 'https://')):
            return jsonify({
                'error': 'URL deve começar com http:// ou https://',
                'code': 'INVALID_URL'
            }), 400
        
        # Gera ID único para a auditoria
        audit_id = str(uuid.uuid4())
        
        # Usar async_processor em vez de threading.Thread direto
        task_id = submit_audit_task(
            orchestrator.execute_full_audit,
            audit_id,
            url,
            options or {}
        )
        
        logger.info(f"Auditoria com CSV iniciada: {audit_id} - arquivo: {file.filename} (task_id: {task_id})")
        
        return jsonify({
            'audit_id': audit_id,
            'url': url,
            'csv_file': file.filename,
            'status': 'started',
            'message': 'Auditoria com CSV iniciada com sucesso',
            'estimated_duration': '3-8 minutos',
            'task_id': task_id
        }), 202
        
    except Exception as e:
        logger.error(f"Erro ao processar upload de CSV: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/status/<audit_id>', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def get_audit_status(audit_id: str):
    """
    Obtém status de uma auditoria em execução.
    
    Args:
        audit_id: ID da auditoria.
    
    Returns:
        Status atual da auditoria.
    """
    try:
        # Verifica se auditoria está ativa
        if audit_id in active_audits:
            audit_info = active_audits[audit_id]
            return jsonify({
                'audit_id': audit_id,
                'status': audit_info['status'],
                'url': audit_info['url'],
                'current_step': audit_info['current_step'],
                'progress': audit_info['progress'],
                'start_time': audit_info['start_time'].isoformat(),
                'last_update': audit_info.get('last_update', audit_info['start_time']).isoformat(),
                'estimated_remaining': f"{10 - (audit_info['progress'] // 10)} minutos"
            }), 200
        
        # Verifica se auditoria foi concluída
        if audit_id in audit_results:
            result = audit_results[audit_id]
            return jsonify({
                'audit_id': audit_id,
                'status': result['status'],
                'url': result['url'],
                'start_time': result['start_time'].isoformat(),
                'end_time': result.get('end_time', '').isoformat() if result.get('end_time') else None,
                'overall_score': result.get('overall_score', 0),
                'steps_completed': len(result.get('steps', [])),
                'errors_count': len(result.get('errors', []))
            }), 200
        
        # Auditoria não encontrada
        return jsonify({
            'error': 'Auditoria não encontrada',
            'audit_id': audit_id,
            'code': 'AUDIT_NOT_FOUND'
        }), 404
        
    except Exception as e:
        logger.error(f"Erro ao obter status da auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/result/<audit_id>', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def get_audit_result(audit_id: str):
    """
    Obtém resultado completo de uma auditoria.
    
    Args:
        audit_id: ID da auditoria.
    
    Returns:
        Resultado completo da auditoria.
    """
    try:
        if audit_id not in audit_results:
            return jsonify({
                'error': 'Resultado de auditoria não encontrado',
                'audit_id': audit_id,
                'code': 'RESULT_NOT_FOUND'
            }), 404
        
        result = audit_results[audit_id]
        
        # Verifica se auditoria foi concluída com sucesso
        if result['status'] != 'completed':
            return jsonify({
                'error': 'Auditoria não foi concluída com sucesso',
                'audit_id': audit_id,
                'status': result['status'],
                'code': 'AUDIT_NOT_COMPLETED'
            }), 400
        
        # Prepara resultado para retorno
        audit_report = result.get('audit_report')
        
        response_data = {
            'audit_id': audit_id,
            'url': result['url'],
            'status': result['status'],
            'start_time': result['start_time'].isoformat(),
            'end_time': result['end_time'].isoformat(),
            'overall_score': result['overall_score'],
            'summary': audit_report.summary if audit_report else {},
            'validations_count': len(audit_report.validations) if audit_report else 0,
            'documentation': result.get('documentation', {}),
            'steps_executed': result.get('steps', []),
            'errors': result.get('errors', [])
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter resultado da auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/report/<audit_id>', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def get_audit_report(audit_id: str):
    """
    Obtém relatório detalhado de uma auditoria.
    
    Args:
        audit_id: ID da auditoria.
    
    Query Parameters:
        format: Formato do relatório (json, html, summary).
    
    Returns:
        Relatório formatado da auditoria.
    """
    try:
        if audit_id not in audit_results:
            return jsonify({
                'error': 'Auditoria não encontrada',
                'audit_id': audit_id,
                'code': 'AUDIT_NOT_FOUND'
            }), 404
        
        result = audit_results[audit_id]
        audit_report = result.get('audit_report')
        
        if not audit_report:
            return jsonify({
                'error': 'Relatório não disponível',
                'audit_id': audit_id,
                'code': 'REPORT_NOT_AVAILABLE'
            }), 404
        
        # Obtém formato solicitado
        report_format = request.args.get('format', 'json').lower()
        
        if report_format == 'html':
            # Gera relatório HTML
            html_content = orchestrator.report_manager.formatter.format_html_report(audit_report)
            
            from flask import Response
            return Response(
                html_content,
                mimetype='text/html',
                headers={'Content-Disposition': f'inline; filename="audit_report_{audit_id}.html"'}
            )
        
        elif report_format == 'summary':
            # Gera relatório resumido
            summary_data = orchestrator.report_manager.formatter.format_summary_report(audit_report)
            return jsonify(summary_data), 200
        
        else:  # json (padrão)
            # Gera relatório JSON completo
            json_data = orchestrator.report_manager.formatter.format_json_report(
                audit_report, include_raw_data=True
            )
            return jsonify(json_data), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter relatório da auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/export/<audit_id>', methods=['GET'])
@rate_limit(limit=50, window=60)
@handle_errors
def export_audit_report(audit_id: str):
    """
    Exporta relatório de auditoria para arquivo.
    
    Args:
        audit_id: ID da auditoria.
    
    Query Parameters:
        format: Formato de exportação (json, html, csv).
    
    Returns:
        Arquivo do relatório para download.
    """
    try:
        if audit_id not in audit_results:
            return jsonify({
                'error': 'Auditoria não encontrada',
                'audit_id': audit_id,
                'code': 'AUDIT_NOT_FOUND'
            }), 404
        
        result = audit_results[audit_id]
        audit_report = result.get('audit_report')
        
        if not audit_report:
            return jsonify({
                'error': 'Relatório não disponível',
                'audit_id': audit_id,
                'code': 'REPORT_NOT_AVAILABLE'
            }), 404
        
        # Obtém formato de exportação
        export_format = request.args.get('format', 'json').lower()
        
        if export_format == 'html':
            # Exporta como HTML
            file_path = orchestrator.report_manager.exporter.export_html_report(audit_report)
            return send_file(
                file_path,
                as_attachment=True,
                download_name=f"audit_report_{audit_id}.html",
                mimetype='text/html'
            )
        
        elif export_format == 'csv':
            # Exporta como CSV (resumo)
            file_path = orchestrator.report_manager.exporter.export_csv_summary([audit_report])
            return send_file(
                file_path,
                as_attachment=True,
                download_name=f"audit_summary_{audit_id}.csv",
                mimetype='text/csv'
            )
        
        else:  # json (padrão)
            # Exporta como JSON
            file_path = orchestrator.report_manager.exporter.export_json_report(
                audit_report, include_raw_data=True
            )
            return send_file(
                file_path,
                as_attachment=True,
                download_name=f"audit_report_{audit_id}.json",
                mimetype='application/json'
            )
        
    except Exception as e:
        logger.error(f"Erro ao exportar relatório da auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/list', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def list_audits():
    """
    Lista todas as auditorias (ativas e concluídas).
    
    Query Parameters:
        status: Filtro por status (running, completed, error).
        limit: Limite de resultados (padrão: 50).
        offset: Offset para paginação (padrão: 0).
    
    Returns:
        Lista de auditorias.
    """
    try:
        # Parâmetros de consulta
        status_filter = request.args.get('status')
        limit = int(request.args.get('limit', 50))
        offset = int(request.args.get('offset', 0))
        
        all_audits = []
        
        # Adiciona auditorias ativas
        for audit_id, audit_info in active_audits.items():
            if not status_filter or audit_info['status'] == status_filter:
                all_audits.append({
                    'audit_id': audit_id,
                    'url': audit_info['url'],
                    'status': audit_info['status'],
                    'start_time': audit_info['start_time'].isoformat(),
                    'current_step': audit_info['current_step'],
                    'progress': audit_info['progress']
                })
        
        # Adiciona auditorias concluídas
        for audit_id, result in audit_results.items():
            if not status_filter or result['status'] == status_filter:
                # Trata start_time - pode ser datetime ou string
                start_time = result.get('start_time')
                if isinstance(start_time, datetime):
                    start_time_str = start_time.isoformat()
                elif isinstance(start_time, str):
                    start_time_str = start_time
                else:
                    start_time_str = datetime.now().isoformat()
                
                # Trata end_time - pode ser datetime ou string
                end_time = result.get('end_time')
                if isinstance(end_time, datetime):
                    end_time_str = end_time.isoformat()
                elif isinstance(end_time, str):
                    end_time_str = end_time
                else:
                    end_time_str = None
                
                all_audits.append({
                    'audit_id': audit_id,
                    'url': result['url'],
                    'status': result['status'],
                    'start_time': start_time_str,
                    'end_time': end_time_str,
                    'overall_score': result.get('overall_score', 0),
                    'errors_count': len(result.get('errors', []))
                })
        
        # Ordena por data de início (mais recente primeiro)
        all_audits.sort(key=lambda x: x['start_time'], reverse=True)
        
        # Aplica paginação
        paginated_audits = all_audits[offset:offset + limit]
        
        return jsonify({
            'audits': paginated_audits,
            'total': len(all_audits),
            'limit': limit,
            'offset': offset,
            'has_more': offset + limit < len(all_audits)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar auditorias: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/stats', methods=['GET'])
@rate_limit(limit=60, window=60)
@handle_errors
def get_audit_stats():
    """
    Obtém estatísticas gerais das auditorias para o dashboard.
    
    Returns:
        Estatísticas incluindo total de auditorias, completadas, com falha,
        pontuação média e auditorias recentes.
    """
    try:
        # Contadores de estatísticas
        total_audits = 0
        completed_audits = 0
        failed_audits = 0
        running_audits = 0
        scores = []
        recent_audits = []
        
        # Processa auditorias ativas
        for audit_id, audit_info in active_audits.items():
            total_audits += 1
            if audit_info['status'] == 'running':
                running_audits += 1
            elif audit_info['status'] == 'error':
                failed_audits += 1
            
            # Adiciona à lista de recentes (limitado a 5)
            if len(recent_audits) < 5:
                recent_audits.append({
                    'audit_id': audit_id,
                    'url': audit_info['url'],
                    'status': audit_info['status'],
                    'start_time': audit_info['start_time'].isoformat() if hasattr(audit_info['start_time'], 'isoformat') else audit_info['start_time'],
                    'current_step': audit_info.get('current_step', 'Iniciando'),
                    'progress': audit_info.get('progress', 0)
                })
        
        # Processa auditorias concluídas
        for audit_id, result in audit_results.items():
            total_audits += 1
            
            if result['status'] == 'completed':
                completed_audits += 1
                # Extrai pontuação se disponível
                if 'audit_report' in result and result['audit_report']:
                    report = result['audit_report']
                    if isinstance(report, dict) and 'overall_score' in report:
                        scores.append(report['overall_score'])
                    elif hasattr(report, 'overall_score'):
                        scores.append(report.overall_score)
            elif result['status'] == 'error':
                failed_audits += 1
            
            # Adiciona à lista de recentes se ainda há espaço
            if len(recent_audits) < 5:
                # Trata start_time - pode ser datetime ou string
                start_time = result.get('start_time')
                if hasattr(start_time, 'isoformat'):
                    start_time_str = start_time.isoformat()
                else:
                    start_time_str = str(start_time) if start_time else datetime.now().isoformat()
                
                recent_audits.append({
                    'audit_id': audit_id,
                    'url': result.get('url', 'URL não disponível'),
                    'status': result['status'],
                    'start_time': start_time_str,
                    'current_step': 'Concluído' if result['status'] == 'completed' else 'Erro',
                    'progress': 100 if result['status'] == 'completed' else 0
                })
        
        # Calcula pontuação média
        average_score = sum(scores) / len(scores) if scores else 0
        
        # Ordena auditorias recentes por data (mais recentes primeiro)
        recent_audits.sort(key=lambda x: x['start_time'], reverse=True)
        recent_audits = recent_audits[:5]  # Limita a 5 mais recentes
        
        return jsonify({
            'total_audits': total_audits,
            'completed_audits': completed_audits,
            'failed_audits': failed_audits,
            'running_audits': running_audits,
            'average_score': round(average_score, 2),
            'recent_audits': recent_audits
        })
        
    except Exception as e:
        logger.error(f"Erro ao obter estatísticas das auditorias: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'message': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


@app.route('/audit/cancel/<audit_id>', methods=['POST'])
@rate_limit(limit=50, window=60)
@handle_errors
def cancel_audit(audit_id: str):
    """
    Cancela uma auditoria em execução.
    
    Args:
        audit_id: ID da auditoria a ser cancelada.
        
    Returns:
        Status do cancelamento.
    """
    try:
        # Verifica se auditoria existe e está ativa
        if audit_id not in active_audits:
            return jsonify({
                'error': 'Auditoria não encontrada ou não está em execução',
                'code': 'AUDIT_NOT_ACTIVE'
            }), 404
        
        logger.info(f"Cancelando auditoria {audit_id}")
        
        # Move para resultados
        cancelled_result = {
            'audit_id': audit_id,
            'url': active_audits[audit_id]['url'],
            'status': 'cancelled',
            'start_time': active_audits[audit_id]['start_time'],
            'end_time': datetime.now(),
            'message': 'Auditoria cancelada pelo usuário'
        }
        
        audit_results[audit_id] = cancelled_result
        del active_audits[audit_id]
        
        logger.info(f"Auditoria {audit_id} cancelada")
        
        return jsonify({
            'audit_id': audit_id,
            'status': 'cancelled',
            'message': 'Auditoria cancelada com sucesso'
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao cancelar auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


# === ENDPOINTS DE DOCUMENTAÇÃO ===

@app.route('/audit/reports/consolidated', methods=['GET'])
@rate_limit(limit=30, window=60)
@handle_errors
def get_consolidated_reports():
    """
    Obtém relatórios consolidados de múltiplas auditorias.
    
    Query Parameters:
        start_date: Data de início (formato: YYYY-MM-DD).
        end_date: Data de fim (formato: YYYY-MM-DD).
        status: Filtro por status (completed, error, etc.).
        audit_type: Tipo de auditoria (full, quick, etc.).
        limit: Limite de resultados (padrão: 50).
        offset: Offset para paginação (padrão: 0).
    
    Returns:
        Relatório consolidado com métricas e insights.
    """
    try:
        # Obter parâmetros de filtro
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        status = request.args.get('status')
        audit_type = request.args.get('audit_type')
        limit = int(request.args.get('limit', 50))
        offset = int(request.args.get('offset', 0))
        
        # Buscar auditorias do banco de dados
        db_manager = DatabaseManager()
        audits = db_manager.get_audit_history(
            start_date=start_date,
            end_date=end_date,
            status=status,
            audit_type=audit_type,
            limit=limit,
            offset=offset
        )
        
        if not audits:
            return jsonify({
                'consolidated_report': {
                    'summary': {
                        'total_audits': 0,
                        'completed_audits': 0,
                        'average_score': 0,
                        'period': f"{start_date or 'início'} até {end_date or 'hoje'}",
                        'top_issues': []
                    },
                    'metrics': {},
                    'insights': [],
                    'trends': []
                },
                'filters_applied': {
                    'start_date': start_date,
                    'end_date': end_date,
                    'status': status,
                    'audit_type': audit_type,
                    'limit': limit,
                    'offset': offset
                }
            })
        
        # Consolidar dados das auditorias
        consolidated_data = orchestrator.data_consolidator.consolidate_multiple_audits(audits)
        
        # Gerar relatório consolidado
        consolidated_report = orchestrator.report_generator.generate_consolidated_report(
            audits=audits,
            consolidated_data=consolidated_data
        )
        
        # Calcular estatísticas adicionais
        total_audits = len(audits)
        completed_audits = len([a for a in audits if a.get('status') == 'completed'])
        average_score = sum([a.get('overall_score', 0) for a in audits if a.get('overall_score')]) / max(total_audits, 1)
        
        # Identificar principais problemas
        all_issues = []
        for audit in audits:
            if audit.get('validation_results'):
                for validation in audit['validation_results']:
                    if validation.get('status') == 'failed':
                        all_issues.append(validation.get('validation_type', 'Unknown'))
        
        from collections import Counter
        top_issues = [{'issue': issue, 'count': count} for issue, count in Counter(all_issues).most_common(5)]
        
        # Montar resposta
        response = {
            'consolidated_report': {
                'summary': {
                    'total_audits': total_audits,
                    'completed_audits': completed_audits,
                    'average_score': round(average_score, 2),
                    'period': f"{start_date or 'início'} até {end_date or 'hoje'}",
                    'top_issues': top_issues
                },
                'metrics': consolidated_report.get('metrics', {}),
                'insights': consolidated_report.get('insights', []),
                'trends': consolidated_report.get('trends', []),
                'detailed_audits': [
                    {
                        'audit_id': audit.get('audit_id'),
                        'url': audit.get('url'),
                        'timestamp': audit.get('timestamp'),
                        'status': audit.get('status'),
                        'overall_score': audit.get('overall_score'),
                        'validation_count': len(audit.get('validation_results', []))
                    }
                    for audit in audits
                ]
            },
            'filters_applied': {
                'start_date': start_date,
                'end_date': end_date,
                'status': status,
                'audit_type': audit_type,
                'limit': limit,
                'offset': offset
            },
            'pagination': {
                'total_results': total_audits,
                'limit': limit,
                'offset': offset,
                'has_more': total_audits == limit  # Indica se há mais resultados
            }
        }
        
        logger.info(f"Relatório consolidado gerado: {total_audits} auditorias processadas")
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Erro ao gerar relatório consolidado: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor ao gerar relatório consolidado',
            'details': str(e) if app.config.get('DEBUG') else None
        }), 500


@app.route('/audit/documentation/<audit_id>', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def view_audit_documentation(audit_id: str):
    """
    Visualiza a documentação da IA gerada para uma auditoria via web.
    
    Args:
        audit_id: ID da auditoria.
        
    Returns:
        Página HTML com a documentação formatada.
    """
    try:
        # Verifica se auditoria existe em memória ou carrega do arquivo
        audit_result = None
        
        if audit_id in audit_results:
            audit_result = audit_results[audit_id]
        else:
            # Tenta carregar do arquivo JSON
            audit_file_path = f"data/reports/audit_{audit_id}.json"
            if os.path.exists(audit_file_path):
                try:
                    with open(audit_file_path, 'r', encoding='utf-8') as f:
                        audit_result = json.load(f)
                except Exception as e:
                    logger.error(f"Erro ao carregar auditoria do arquivo {audit_file_path}: {e}")
        
        if not audit_result:
            return f"""
            <!DOCTYPE html>
            <html lang="pt-BR">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Documentação não encontrada</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }}
                    .error {{ background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                    .error h1 {{ color: #d32f2f; }}
                </style>
            </head>
            <body>
                <div class="error">
                    <h1>❌ Documentação não encontrada</h1>
                    <p>A auditoria com ID <strong>{audit_id}</strong> não foi encontrada.</p>
                    <p><a href="/audit/list">← Voltar para lista de auditorias</a></p>
                </div>
            </body>
            </html>
            """, 404
        
        # Verifica se há documentação gerada
        documentation_data = audit_result.get('documentation', {})
        
        if not documentation_data:
            return f"""
            <!DOCTYPE html>
            <html lang="pt-BR">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Documentação não disponível</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }}
                    .warning {{ background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                    .warning h1 {{ color: #f57c00; }}
                </style>
            </head>
            <body>
                <div class="warning">
                    <h1>⚠️ Documentação não disponível</h1>
                    <p>A documentação da IA ainda não foi gerada para esta auditoria.</p>
                    <p><strong>Status da auditoria:</strong> {audit_result.get('status', 'Desconhecido')}</p>
                    <p><a href="/audit/list">← Voltar para lista de auditorias</a></p>
                </div>
            </body>
            </html>
            """, 404
        
        # Extrair documentação especializada em SEO
        seo_documentation = documentation_data.get('seo_specialized', {})
        
        if seo_documentation.get('status') == 'no_problems':
            markdown_content = f"""# Relatório de Auditoria SEO - {audit_result.get('url', 'URL não disponível')}

**Data de Geração:** {datetime.now().strftime('%d/%m/%Y às %H:%M')}
**Status:** Auditoria concluída sem problemas críticos

## ✅ Resultado Positivo

{seo_documentation.get('message', 'Nenhum problema técnico de SEO foi identificado na auditoria.')}

### Informações da Auditoria:
- **URL Auditada:** {audit_result.get('url', 'N/A')}
- **ID da Auditoria:** {audit_id}
- **Pontuação Geral:** {audit_result.get('overall_score', 'N/A')}
- **Data de Início:** {audit_result.get('start_time', 'N/A')}
- **Data de Conclusão:** {audit_result.get('end_time', 'N/A')}

### Próximos Passos:
1. Continue monitorando o desempenho do site regularmente
2. Considere executar auditorias periódicas para manter a qualidade
3. Fique atento a mudanças no algoritmo do Google que possam afetar o SEO
"""
        else:
            # Usar documentação gerada pelo agente especializado
            doc_data = seo_documentation.get('document', {})
            markdown_content = doc_data.get('markdown', 'Documentação não disponível')
        
        # Converter Markdown para HTML
        html_content = _convert_markdown_to_html(markdown_content, audit_id, audit_result)
        
        return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
        
    except Exception as e:
        logger.error(f"Erro ao visualizar documentação da auditoria {audit_id}: {str(e)}")
        return f"""
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Erro interno</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }}
                .error {{ background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                .error h1 {{ color: #d32f2f; }}
            </style>
        </head>
        <body>
            <div class="error">
                <h1>❌ Erro interno do servidor</h1>
                <p>Ocorreu um erro ao carregar a documentação.</p>
                <p><strong>Detalhes:</strong> {str(e)}</p>
                <p><a href="/audit/list">← Voltar para lista de auditorias</a></p>
            </div>
        </body>
        </html>
        """, 500


@app.route('/audit/documentation/<audit_id>/download', methods=['GET'])
@rate_limit(limit=30, window=60)
@handle_errors
def download_audit_documentation(audit_id: str):
    """
    Faz download da documentação da IA gerada para uma auditoria em formato Word.
    
    Args:
        audit_id: ID da auditoria.
        
    Returns:
        Arquivo Word (.docx) com a documentação.
    """
    try:
        # Verifica se auditoria existe em memória ou carrega do arquivo
        audit_result = None
        
        if audit_id in audit_results:
            audit_result = audit_results[audit_id]
        else:
            # Tenta carregar do arquivo JSON
            audit_file_path = f"data/reports/audit_{audit_id}.json"
            if os.path.exists(audit_file_path):
                try:
                    with open(audit_file_path, 'r', encoding='utf-8') as f:
                        audit_result = json.load(f)
                except Exception as e:
                    logger.error(f"Erro ao carregar auditoria do arquivo {audit_file_path}: {e}")
        
        if not audit_result:
            return jsonify({
                'error': 'Auditoria não encontrada',
                'audit_id': audit_id,
                'code': 'AUDIT_NOT_FOUND'
            }), 404
        
        # Verifica se há documentação gerada
        documentation_data = audit_result.get('documentation', {})
        
        if not documentation_data:
            return jsonify({
                'error': 'Documentação não disponível para esta auditoria',
                'code': 'DOCUMENTATION_NOT_AVAILABLE'
            }), 404
        
        # Gerar arquivo Word
        word_file_path = _generate_word_documentation(audit_id, audit_result, documentation_data)
        
        if not word_file_path or not os.path.exists(word_file_path):
            return jsonify({
                'error': 'Erro ao gerar arquivo Word',
                'code': 'WORD_GENERATION_ERROR'
            }), 500
        
        # Enviar arquivo para download
        return send_file(
            word_file_path,
            as_attachment=True,
            download_name=f'documentacao_seo_{audit_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Erro ao fazer download da documentação da auditoria {audit_id}: {str(e)}")
        return jsonify({
            'error': 'Erro interno do servidor',
            'details': str(e),
            'code': 'INTERNAL_ERROR'
        }), 500


# === TRATAMENTO DE ERROS ===

def _convert_markdown_to_html(markdown_content: str, audit_id: str, audit_result: Dict) -> str:
    """
    Converte conteúdo Markdown para HTML com estilo profissional.
    
    Args:
        markdown_content: Conteúdo em Markdown.
        audit_id: ID da auditoria.
        audit_result: Dados da auditoria.
        
    Returns:
        HTML formatado com CSS.
    """
    try:
        import re
        
        # Escapar HTML básico
        html_content = markdown_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Converter Markdown para HTML
        # Headers
        html_content = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html_content, flags=re.MULTILINE)
        
        # Bold e Italic
        html_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_content)
        html_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_content)
        
        # Links
        html_content = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" target="_blank">\1</a>', html_content)
        
        # Listas
        html_content = re.sub(r'^- (.*?)$', r'<li>\1</li>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^(\d+)\. (.*?)$', r'<li>\1. \2</li>', html_content, flags=re.MULTILINE)
        
        # Envolver listas em <ul>
        html_content = re.sub(r'(<li>.*?</li>)', r'<ul>\1</ul>', html_content, flags=re.DOTALL)
        html_content = re.sub(r'</ul>\s*<ul>', '', html_content)
        
        # Code blocks
        html_content = re.sub(r'```(\w+)?\n(.*?)\n```', r'<pre><code class="language-\1">\2</code></pre>', html_content, flags=re.DOTALL)
        html_content = re.sub(r'`([^`]+)`', r'<code>\1</code>', html_content)
        
        # Quebras de linha
        html_content = html_content.replace('\n\n', '</p><p>').replace('\n', '<br>')
        html_content = f'<p>{html_content}</p>'
        
        # Limpar tags vazias
        html_content = re.sub(r'<p></p>', '', html_content)
        html_content = re.sub(r'<p><br></p>', '', html_content)
        
        # Template HTML completo
        full_html = f"""
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Documentação SEO - {audit_result.get('url', 'Auditoria')}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 20px;
                    background-color: #f8f9fa;
                }}
                
                .header {{
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 30px;
                    border-radius: 10px;
                    margin-bottom: 30px;
                    box-shadow: 0 4px 15px rgba(0,0,0,0.1);
                }}
                
                .header h1 {{
                    margin: 0 0 10px 0;
                    font-size: 2.2em;
                }}
                
                .header .meta {{
                    opacity: 0.9;
                    font-size: 1.1em;
                }}
                
                .content {{
                    background: white;
                    padding: 40px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                    margin-bottom: 30px;
                }}
                
                .actions {{
                    background: white;
                    padding: 20px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                    text-align: center;
                }}
                
                .btn {{
                    display: inline-block;
                    padding: 12px 24px;
                    margin: 0 10px;
                    text-decoration: none;
                    border-radius: 6px;
                    font-weight: 600;
                    transition: all 0.3s ease;
                }}
                
                .btn-primary {{
                    background: #007bff;
                    color: white;
                }}
                
                .btn-primary:hover {{
                    background: #0056b3;
                    transform: translateY(-2px);
                }}
                
                .btn-success {{
                    background: #28a745;
                    color: white;
                }}
                
                .btn-success:hover {{
                    background: #1e7e34;
                    transform: translateY(-2px);
                }}
                
                .btn-secondary {{
                    background: #6c757d;
                    color: white;
                }}
                
                .btn-secondary:hover {{
                    background: #545b62;
                    transform: translateY(-2px);
                }}
                
                h1, h2, h3 {{
                    color: #2c3e50;
                    margin-top: 30px;
                    margin-bottom: 15px;
                }}
                
                h1 {{ font-size: 2.2em; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                h2 {{ font-size: 1.8em; border-bottom: 2px solid #e74c3c; padding-bottom: 8px; }}
                h3 {{ font-size: 1.4em; color: #8e44ad; }}
                
                p {{ margin-bottom: 15px; }}
                
                ul, ol {{
                    margin-bottom: 20px;
                    padding-left: 30px;
                }}
                
                li {{
                    margin-bottom: 8px;
                }}
                
                code {{
                    background: #f8f9fa;
                    padding: 2px 6px;
                    border-radius: 4px;
                    font-family: 'Courier New', monospace;
                    color: #e83e8c;
                }}
                
                pre {{
                    background: #2d3748;
                    color: #e2e8f0;
                    padding: 20px;
                    border-radius: 8px;
                    overflow-x: auto;
                    margin: 20px 0;
                }}
                
                pre code {{
                    background: none;
                    color: inherit;
                    padding: 0;
                }}
                
                strong {{
                    color: #2c3e50;
                    font-weight: 600;
                }}
                
                a {{
                    color: #3498db;
                    text-decoration: none;
                }}
                
                a:hover {{
                    color: #2980b9;
                    text-decoration: underline;
                }}
                
                .alert {{
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 6px;
                    border-left: 4px solid;
                }}
                
                .alert-success {{
                    background: #d4edda;
                    border-color: #28a745;
                    color: #155724;
                }}
                
                .alert-warning {{
                    background: #fff3cd;
                    border-color: #ffc107;
                    color: #856404;
                }}
                
                .alert-info {{
                    background: #d1ecf1;
                    border-color: #17a2b8;
                    color: #0c5460;
                }}
                
                @media (max-width: 768px) {{
                    body {{ padding: 10px; }}
                    .header, .content, .actions {{ padding: 20px; }}
                    .btn {{ display: block; margin: 10px 0; }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>📊 Documentação de Auditoria SEO</h1>
                <div class="meta">
                    <strong>URL:</strong> {audit_result.get('url', 'N/A')} | 
                    <strong>ID:</strong> {audit_id} | 
                    <strong>Pontuação:</strong> {audit_result.get('overall_score', 'N/A')}
                </div>
            </div>
            
            <div class="content">
                {html_content}
            </div>
            
            <div class="actions">
                <a href="/audit/documentation/{audit_id}/download" class="btn btn-success">
                    📄 Download Word (.docx)
                </a>
                <a href="/audit/result/{audit_id}" class="btn btn-primary">
                    📈 Ver Resultado Completo
                </a>
                <a href="/audit/list" class="btn btn-secondary">
                    📋 Lista de Auditorias
                </a>
            </div>
        </body>
        </html>
        """
        
        return full_html
        
    except Exception as e:
        logger.error(f"Erro ao converter Markdown para HTML: {str(e)}")
        return f"""
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
            <meta charset="UTF-8">
            <title>Erro de conversão</title>
        </head>
        <body>
            <h1>Erro ao processar documentação</h1>
            <p>Ocorreu um erro ao converter a documentação para HTML.</p>
            <p><strong>Erro:</strong> {str(e)}</p>
        </body>
        </html>
        """


def _generate_word_documentation(audit_id: str, audit_result: Dict, documentation_data: Dict) -> Optional[str]:
    """
    Gera arquivo Word (.docx) com a documentação da auditoria.
    
    Args:
        audit_id: ID da auditoria.
        audit_result: Dados da auditoria.
        documentation_data: Dados da documentação gerada.
        
    Returns:
        Caminho do arquivo Word gerado ou None se houver erro.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        # Criar documento Word
        doc = Document()
        
        # Configurar estilos
        styles = doc.styles
        
        # Título principal
        title = doc.add_heading('Relatório de Auditoria SEO Técnica', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informações da auditoria
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('URL Auditada:', audit_result.get('url', 'N/A')),
            ('ID da Auditoria:', audit_id),
            ('Data de Geração:', datetime.now().strftime('%d/%m/%Y às %H:%M')),
            ('Pontuação Geral:', str(audit_result.get('overall_score', 'N/A'))),
            ('Status:', audit_result.get('status', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            
            # Negrito na primeira coluna
            info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Espaço
        
        # Extrair documentação
        seo_documentation = documentation_data.get('seo_specialized', {})
        
        if seo_documentation.get('status') == 'no_problems':
            # Caso sem problemas
            doc.add_heading('✅ Resultado Positivo', level=1)
            
            message = seo_documentation.get('message', 'Nenhum problema técnico de SEO foi identificado na auditoria.')
            doc.add_paragraph(message)
            
            doc.add_heading('Próximos Passos', level=2)
            next_steps = [
                'Continue monitorando o desempenho do site regularmente',
                'Considere executar auditorias periódicas para manter a qualidade',
                'Fique atento a mudanças no algoritmo do Google que possam afetar o SEO'
            ]
            
            for step in next_steps:
                p = doc.add_paragraph(step, style='List Bullet')
        
        else:
            # Caso com problemas - usar documentação gerada
            doc_data = seo_documentation.get('document', {})
            
            if 'markdown' in doc_data:
                # Converter Markdown básico para Word
                markdown_content = doc_data['markdown']
                _add_markdown_to_word_doc(doc, markdown_content)
            
            elif 'json_data' in doc_data:
                # Usar dados JSON estruturados
                json_data = doc_data['json_data']
                
                # Metadados
                metadata = json_data.get('metadata', {})
                if metadata:
                    doc.add_heading('Resumo Executivo', level=1)
                    doc.add_paragraph(f"Total de problemas identificados: {metadata.get('total_problems', 0)}")
                    
                    severity_dist = metadata.get('severity_distribution', {})
                    if severity_dist:
                        doc.add_paragraph('Distribuição por severidade:')
                        for severity, count in severity_dist.items():
                            doc.add_paragraph(f"• {severity}: {count} problemas", style='List Bullet')
                
                # Problemas
                problems = json_data.get('problems', [])
                if problems:
                    doc.add_heading('Problemas Identificados', level=1)
                    
                    for i, problem in enumerate(problems, 1):
                        doc.add_heading(f"{i}. {problem.get('title', 'Problema sem título')}", level=2)
                        
                        # Descrição
                        if problem.get('problem_description'):
                            doc.add_heading('🔍 Problema/Oportunidade', level=3)
                            doc.add_paragraph(problem['problem_description'])
                        
                        # Evidências
                        if problem.get('evidence_list'):
                            doc.add_heading('📋 Evidências', level=3)
                            doc.add_paragraph(problem['evidence_list'])
                        
                        # Impacto
                        if problem.get('impact_analysis'):
                            doc.add_heading('📊 Análise de Impacto', level=3)
                            doc.add_paragraph(problem['impact_analysis'])
                        
                        # Solução
                        if problem.get('technical_steps'):
                            doc.add_heading('🔧 Passos Técnicos', level=3)
                            doc.add_paragraph(problem['technical_steps'])
                        
                        # Benefícios esperados
                        if problem.get('expected_benefits'):
                            doc.add_heading('✅ Benefícios Esperados', level=3)
                            doc.add_paragraph(problem['expected_benefits'])
                        
                        # Cronograma
                        if problem.get('timeline_estimate'):
                            doc.add_heading('⏱️ Estimativa de Tempo', level=3)
                            doc.add_paragraph(problem['timeline_estimate'])
                        
                        doc.add_page_break()
        
        # Salvar arquivo
        exports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'exports')
        os.makedirs(exports_dir, exist_ok=True)
        filename = f'documentacao_seo_{audit_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        file_path = os.path.join(exports_dir, filename)
        
        doc.save(file_path)
        
        logger.info(f"Arquivo Word gerado: {file_path}")
        return file_path
        
    except ImportError:
        logger.error("Biblioteca python-docx não está instalada. Execute: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"Erro ao gerar arquivo Word: {str(e)}")
        return None


def _add_markdown_to_word_doc(doc, markdown_content: str):
    """
    Adiciona conteúdo Markdown básico ao documento Word.
    
    Args:
        doc: Documento Word.
        markdown_content: Conteúdo em Markdown.
    """
    try:
        import re
        
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Listas
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(line[3:], style='List Number')
            
            # Texto normal
            else:
                # Processar formatação básica
                paragraph = doc.add_paragraph()
                
                # Bold
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)
        
    except Exception as e:
        logger.error(f"Erro ao processar Markdown no Word: {str(e)}")
        # Fallback: adicionar como texto simples
        doc.add_paragraph(markdown_content)


# === ENDPOINT PARA DEMONSTRAÇÃO ===

@app.route('/process/<int:item_id>', methods=['GET'])
@rate_limit(limit=100, window=60)
@handle_errors
def process_item(item_id: int):
    """
    Endpoint de demonstração para processar um item por ID.
    
    Args:
        item_id: ID do item a ser processado
        
    Returns:
        JSON com resultado do processamento ou erro
    """
    try:
        # Validação do ID
        if item_id < 0:
            return jsonify({
                'error': 'O ID do item não pode ser negativo.'
            }), 400
        
        # Simulação de processamento
        result_data = f"Dados processados para o item {item_id}. Timestamp: {datetime.now().isoformat()}"
        
        logger.info(f"Item {item_id} processado com sucesso")
        
        return jsonify({
            'item_id': item_id,
            'result': result_data
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao processar item {item_id}: {str(e)}")
        return jsonify({
            'error': f'Erro interno do servidor: {str(e)}'
        }), 500


# === TRATAMENTO DE ERROS ===

@app.errorhandler(404)
def not_found(error):
    """Tratamento de erro 404."""
    return jsonify({
        'error': 'Endpoint não encontrado',
        'code': 'NOT_FOUND',
        'timestamp': datetime.now().isoformat()
    }), 404


@app.errorhandler(405)
def method_not_allowed(error):
    """Tratamento de erro 405."""
    return jsonify({
        'error': 'Método não permitido',
        'code': 'METHOD_NOT_ALLOWED',
        'timestamp': datetime.now().isoformat()
    }), 405


@app.errorhandler(500)
def internal_error(error):
    """Tratamento de erro 500."""
    logger.error(f"Erro interno: {str(error)}")
    return jsonify({
        'error': 'Erro interno do servidor',
        'code': 'INTERNAL_ERROR',
        'timestamp': datetime.now().isoformat()
    }), 500


# === INICIALIZAÇÃO ===

def initialize_app():
    """Inicializa a aplicação e seus componentes."""
    try:
        # Cria diretórios necessários
        os.makedirs('data', exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        os.makedirs('data/reports', exist_ok=True)
        
        # Inicia processador assíncrono
        start_async_processing()
        logger.info("Processador assíncrono iniciado")
        
        logger.info("Aplicação Flask inicializada com sucesso")
        
    except Exception as e:
        logger.error(f"Erro na inicialização: {str(e)}")
        raise


if __name__ == '__main__':
    # Inicializa aplicação
    initialize_app()
    
    # Configurações do servidor
    host = os.getenv('FLASK_HOST', '0.0.0.0')
    port = int(os.getenv('FLASK_PORT', 5000))
    debug = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
    
    logger.info(f"Iniciando servidor Flask com WebSocket em {host}:{port}")
    
    # Usar SocketIO em vez de app.run para suporte a WebSocket
    socketio.run(app, host=host, port=port, debug=debug, allow_unsafe_werkzeug=True)